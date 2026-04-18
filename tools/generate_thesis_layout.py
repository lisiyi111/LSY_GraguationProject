from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


FIGURE_SECTION = """
### 2.6 论文配图设计说明

为使论文图文对应关系更加清晰，结合当前项目代码结构、系统交互流程以及设计图中的协议说明，本文建议在论文中补充系统结构图、系统流程图、系统模块图和系统用例图四类核心图示。四张图的作用并不相同：系统结构图强调“上位机与下位机的层次关系”，系统流程图强调“用户操作和数据流转过程”，系统模块图强调“软件内部模块组织”，系统用例图强调“用户和系统之间的功能交互边界”。下面对各图的绘制内容和正文配套说明进行统一设计。

#### 2.6.1 系统结构图说明

建议将该图命名为“图2-1 彩灯控制系统总体结构图”。图中可分为五个层次，自上而下依次为用户层、上位机软件层、通信传输层、主下位机层和从机及彩灯执行层。用户层主要体现操作人员；上位机软件层包括三维可视化界面、场景管理模块、BIN文件生成模块和串口控制模块；通信传输层体现串口、无线链路或RS-485总线；主下位机层体现协议解析、命令转发和状态回传；执行层体现31组彩灯节点及摄像头状态。

在正文中可配套写为：系统总体结构采用“上位机统一配置、主下位机分组转发、从机节点具体执行”的分层控制方式。上位机APP负责完成参数编辑、场景组织、协议封装和命令下发；主下位机负责接收上位机数据并向组内从机转发；从机节点根据接收到的RGBW参数驱动彩灯工作，并在线路检查时将在线状态逐级回传。

#### 2.6.2 系统流程图说明

建议将该图命名为“图2-2 上位机软件工作流程图”。流程起点为软件启动，随后依次进入模型显示、面选择、单灯选中、RGBW参数编辑、场景保存或插入、BIN文件生成、串口发送、运行命令下发和反馈解析等步骤。流程图中应体现至少两个判断分支：一是“是否保存当前场景”，二是“是否进入串口模式进行联调”。如果选择继续编辑，则返回模型界面；如果选择发送，则进入串口发送和日志反馈流程。

在正文中可配套写为：系统工作流程以“编辑场景”和“联机下发”两条主线展开。用户首先在三维模型中完成灯具参数配置，随后将当前状态写入场景列表；当场景编排完成后，系统按照既定顺序生成二进制文件，并在串口模式下向下位机发送数据或控制命令。若收到运行成功、清除成功或线路检查反馈，则系统进一步在日志区输出解析结果。

#### 2.6.3 系统模块图说明

建议将该图命名为“图4-1 上位机软件功能模块图”。图中可分为界面交互模块、灯具对象管理模块、面切换模块、场景管理模块、BIN文件读写模块、串口通信模块和日志反馈模块。若与代码对应，可进一步在图中标注主要脚本名称：`Lamp.cs`、`LampManager.cs`、`FaceManager.cs`、`FaceLightManager.cs`、`SceneUIController.cs`、`SceneBinSaver.cs`、`SceneBinLoader.cs`、`SerialManager.cs`、`SerialUIController.cs`、`UILogger.cs`等。

在正文中可配套写为：软件采用模块化设计思路，将彩灯对象建模、参数编辑、场景组织、文件封装和串口调试分别划归到不同模块中。模块之间通过对象引用、事件回调和统一的数据结构进行协作，既降低了耦合度，又有利于后期维护和扩展。其中，场景管理模块与文件读写模块共同构成数据核心，串口通信模块与日志反馈模块共同构成联调核心。

#### 2.6.4 系统用例图说明

建议将该图命名为“图2-3 上位机软件用例图”。图中的主要参与者为“操作人员”，核心用例包括旋转模型、选择面、选中彩灯、编辑RGBW参数、设置摄像头状态、添加场景、插入场景、加载BIN文件、生成BIN文件、打开串口、发送BIN文件、发送运行命令、发送清除命令、发送线路检查命令和查看反馈日志。若希望进一步细化，可把“编辑当前场景”作为“选择场景”的扩展用例，把“查看异常灯具信息”作为“发送线路检查命令”的结果用例。

在正文中可配套写为：从用户视角看，上位机系统既承担离线配置任务，也承担在线联调任务。前者主要包括模型观察、参数编辑和场景组织，后者主要包括串口连接、文件下发、命令控制和反馈查看。用例图能够从需求层面对系统功能边界进行归纳，为后续模块划分与测试设计提供依据。
"""


FRONT_MATTER = """华中科技大学本科生毕业设计（论文）

# 基于编码彩灯的上位机APP设计

学院：电子信息与通信学院  
专业：通信工程  
班级：电信2201班  
学生姓名：李思懿  
学号：U202213929  
指导教师：龚军  

\f

## 摘要
"""


TOC_BLOCK = """
\f

## 目录

说明：本节在Word中使用“引用-目录-自动目录”生成，提交终稿前统一更新页码。

\f
"""


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r = run._element.rPr
    r.rFonts.set(qn("w:eastAsia"), font_name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name, font_name, size in [
        ("Title", "黑体", 18),
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
        ("Heading 4", "黑体", 12),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_east_asia_font(run, "宋体")
    run.font.size = Pt(10.5)


def build_markdown(src: str) -> str:
    abstract_pos = src.find("## 摘要")
    if abstract_pos == -1:
        raise ValueError("未找到摘要标题")

    body = src[abstract_pos + len("## 摘要") :].lstrip()
    merged = FRONT_MATTER + body

    lines = merged.splitlines()
    result: list[str] = []
    inserted_toc = False
    inserted_figures = False

    for line in lines:
        if not inserted_toc and line.startswith("## 第一章 "):
            result.extend(TOC_BLOCK.strip("\n").splitlines())
            inserted_toc = True

        if not inserted_figures and line.startswith("## 第三章 "):
            result.extend(FIGURE_SECTION.strip("\n").splitlines())
            result.append("")
            inserted_figures = True

        result.append(line)

    return "\n".join(result) + "\n"


def markdown_to_docx(md_text: str, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])

    lines = md_text.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if line == "\f":
            doc.add_page_break()
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            set_east_asia_font(run, "黑体")
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            run = p.add_run(line[3:].strip())
            set_east_asia_font(run, "黑体")
            continue

        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(line[4:].strip())
            set_east_asia_font(run, "黑体")
            continue

        if line.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            run = p.add_run(line[5:].strip())
            set_east_asia_font(run, "黑体")
            continue

        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(line)
            set_east_asia_font(run, "宋体")
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="Normal")
            run = p.add_run("• " + line[2:])
            set_east_asia_font(run, "宋体")
            continue

        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(line)
        set_east_asia_font(run, "宋体")

    doc.save(output_path)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    src = next(root.glob("*初稿.md"))
    src_text = src.read_text(encoding="utf-8")

    output_md = root / "毕业论文排版稿.md"
    output_docx = root / "毕业论文排版稿.docx"

    final_md = build_markdown(src_text)
    output_md.write_text(final_md, encoding="utf-8")
    markdown_to_docx(final_md, output_docx)

    print(output_md)
    print(output_docx)


if __name__ == "__main__":
    main()
